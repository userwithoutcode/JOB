import docx
import os
from pprint import pprint


def getText(filename):
    """
    Работает с текстом
    """
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText


def open_tables(filename):
    """
       Работает с таблицами
    """
    document = docx.Document(filename)
    data = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    data.append(para.text)
                    # print(para.text)
    return data


def get_list(data, data1):
    lst1 = []
    lst2 = []
    lst3 = []
    lst4 = []
    for items in data:
        if 'От кого' in items:
            x = items
            lst1.append(x.strip('От кого'))
    for number in data1:
        if '#' in number:
            x = number
            lst2.append(x.strip(' .. '))
    for name in data:
        if 'Клиент' in name or 'Узел' in name:
            n = name
            lst3.append(n)
    for adress in data:
        if 'Адрес' in adress:
            x = adress
            lst4.append(x.strip('Адрес:'))
    lstx = zip(lst1, lst2, lst3, lst4)
    return lstx


def get_result_list(datax):
    count = 0
    result_list = []
    for namex, numberx, objectx, adressx in datax:
        count += 1
        x = namex
        y = numberx
        z = objectx
        h = adressx
        result_list.append('{}. {}, {}, {}, {}'.format(count, h, y, x, z))
    return result_list


def get_list_of_words(data):
    lst1 = []
    list_of_words = []
    aprove_words = ['Боровский', 'Стандарчук', 'Агеев', 'Толпекин', 'Медведев', 'Кудрявцев']
    for items in data:
        if 'От кого' in items:
            x = items
            lst1.append(x.strip('От кого'))
    for item in lst1:
        new_list = item.split()
        for word in new_list:
            if word in aprove_words:
                list_of_words.append(word)
    return(list_of_words)


def get_wordfreq(list_of_words):
    '''
    Считаем количество задач, созданные ПМ
    '''
    wordfreq = [list_of_words.count(w) for w in list_of_words]
    draft = list(zip(list_of_words, wordfreq))
    draft_one = set(draft)
    sort_draft = sorted(draft_one, key=lambda x: x[1], reverse=True)
    return sort_draft


def write_to_file(result_list, filename):
    '''
    Записываем все данные в файл
    '''
    with open(filename, 'w', encoding='utf8') as f:
        for line in result_list:
            f.write(line + '\n')
        f.write('\n' + 'Общее колиество задач: ' + str(total_quantity) + '\n')
        for line in sort_draft:
            f.write('\n' + str(line))


if __name__ == '__main__':
    filename = '1.docx'
    data = open_tables(filename)
    data1 = getText(filename)
    datax = get_list(data, data1)
    result_list = get_result_list(datax)
    total_quantity = len(result_list)
    list_of_words = get_list_of_words(data)
    sort_draft = get_wordfreq(list_of_words)
    name, extension = os.path.splitext(filename)
    name_result = name + '.txt'
    write_to_file(result_list, name_result)
